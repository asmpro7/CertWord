# CertWord script
# made by asmpro
# date: 16/4/2023
# TG:@asmprotk

import docx
from docx.shared import Pt, RGBColor
import shutil

cert = docx.Document("cert.docx")
names = docx.Document("names.docx")
ListNames = []
for name in names.paragraphs:
    ListNames.append(name.text)
for name in ListNames:
    shutil.copyfile("cert.docx", f"certs\{name}.docx")
    FName = docx.Document(f"certs\{name}.docx")
    for para in FName.paragraphs:
        if "old" in para.text:
            for run in para.runs:
                if "old" in run.text:
                    run.text = run.text.replace('old', f'{name}')
                    font = run.font
                    font.name = 'Script MT Bold'
                    font.size = Pt(50)
                    font.color.rgb = RGBColor(40, 120, 135)
    FName.save(f"certs\{name}.docx")
